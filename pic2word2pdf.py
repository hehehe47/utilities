#!/usr/bin/env python 
# -*- coding: utf-8 -*- 
# @Time : 2020/3/24 0:02 
# @Author : Patrick 
# @File : pic2word2pdf.py 
# @Software: PyCharm


import os

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches

doc = Document()  # doc对象
images_dir = 'F:\Guitar\喜帖街'
section = doc.sections[-1]
new_width, new_height = section.page_height, section.page_width
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = new_width
section.page_height = new_height
string = '文字内容'
# images = '1.jpg'    # 保存在本地的图片
doc.add_paragraph(string)  # 添加文字
for img in os.listdir(images_dir):
    image_dir = os.path.join(images_dir, img)
    single_img = Image.open(image_dir)
    doc.add_picture(image_dir, width=Inches(1.0))  # 添加图, 设置宽度
doc.save('test.docx')  # 保存路径
